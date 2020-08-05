import sys
import tweepy
import csv
import re
import spacy
from textblob import TextBlob
import matplotlib.pyplot as plt
from firebase import firebase
from docx import Document
from docx.shared import Inches

nlp = spacy.load("C:/Users/SMIT PATEL/.conda/envs/myEnv/Lib/site-packages/en_core_web_lg/en_core_web_lg-2.2.5")
firebase1 = firebase.FirebaseApplication('https://feedback-analysis-a06eb.firebaseio.com/',None)
keys = firebase1.get('/Expert_Name',None)
firebase1 = firebase.FirebaseApplication('https://feedback-analysis-a06eb.firebaseio.com/Expert_Name/',None)
positivefeedback = open('positivefeedback.txt','w')
positivefeedback.write('Positive :\n')
positivefeedback.close()
negativefeedback = open('negativefeedback.txt','w')
negativefeedback.write('negative :\n')
negativefeedback.close()
neutralfeedback = open('neutralfeedback.txt','w')
neutralfeedback.write('neutral :\n')
neutralfeedback.close()
finalfeedback = open('finalfeedback.docx','w')
finalfeedback.write('Final Report :\n')
finalfeedback.close()
positive_list=[]
negative_list=[]
neutral_list=[]
date=''
Expert_Name=''
Subject_Name=''
a=[]
x=0 
for keyID in keys:
            Date = firebase1.get(keyID,'Date')
            print(type(Date))
            print(Date[ 0 : 10 ])
            x=x+1
            for i in Date[ 0 : 10 ]:
               if (i!='-' and x==1):
                  Expert_Name = firebase1.get(keyID,'Expert_Name')
                  Subject_Name = firebase1.get(keyID,'Subject_Name')
                  a.append(i)
            date=a[6]+a[7]+"/"+a[4]+a[5]+"/"+a[0]+a[1]+a[2]+a[3]
            print(date)
class Sentiment_Analysis:
    def fatchdata(self):
        positive = 0
        neutral = 0
        negative = 0
        polarity = 0
        totalfeedback=0
        
        print('\nPlease wait while we are processing your request...')
        self.queation_analysis(1)
        self.queation_analysis(2)
        self.queation_analysis(3)
        self.queation_analysis(4)
        self.queation_analysis(5)
        for keyID in keys:
            totalfeedback += 1
            feedback = firebase1.get(keyID,'Feedback')
            print(type(feedback))
            print(feedback)
            analysis = TextBlob(feedback)
            polarity += analysis.sentiment.polarity  
            if (analysis.sentiment.polarity == 0): 
                neutral += 1
                neutral_list.append(feedback)
                neutralfeedback = open('neutralfeedback.txt','a')
                neutralfeedback.write(feedback)
                neutralfeedback.write('\n')
                neutralfeedback.close()
            elif (analysis.sentiment.polarity > 0.0 and analysis.sentiment.polarity <= 1.0):
                positive += 1
                positive_list.append(feedback)
                positivefeedback = open('positivefeedback.txt','a')
                positivefeedback.write(feedback)
                positivefeedback.write('\n')
                positivefeedback.close() 
            elif (analysis.sentiment.polarity >= -1.0 and analysis.sentiment.polarity < 0.0):
                negative += 1
                negative_list.append(feedback)
                negativefeedback = open('negativefeedback.txt','a')
                negativefeedback.write(feedback)
                negativefeedback.write('\n')
                negativefeedback.close()

        positive = self.percentage(positive, totalfeedback)
        negative = self.percentage(negative, totalfeedback)
        neutral = self.percentage(neutral, totalfeedback)
        polarity = polarity / totalfeedback
        print()
        print("General Report: ")
        if (polarity == 0):
            print("Neutral")
        elif (polarity > 0.0 and polarity <= 1.0):
            print("Positive")
        elif (polarity >= -1.0 and polarity < 0.0):
            print("Negative")
        print()
        print("Detailed Report: ")
        print(str(positive) + "% people thought it was positive")
        print(str(negative) + "% people thought it was negative")
        print(str(neutral) + "% people thought it was neutral")
        print(positive_list)
        print(neutral_list)
        print(negative_list)
        self.removepositive(positive_list)
        self.removenegative(negative_list)
        self.removeneutral(neutral_list)
        self.plotPieChart(positive, negative, neutral, totalfeedback)
        

        finalfeedback = open('finalfeedback.docx','a')
        finalfeedback.write("\n\nPositive Feedback:\n")
        for i in positive_list:
            finalfeedback.write(i)
            finalfeedback.write('\n') 
        finalfeedback.write("\n\nNegative Feedback:\n")
        for i in negative_list:
            finalfeedback.write(i)
            finalfeedback.write('\n')
        finalfeedback.write("\n\nNeutral Feedback:\n")
        for i in neutral_list:
            finalfeedback.write(i)
            finalfeedback.write('\n')
        finalfeedback.write('\n')
        finalfeedback.close()

    def queation_analysis(self,i):
        Excelent=0
        Good=0
        Average=0
        Poor=0
        tf = 0
        for keyID in keys:
            tf += 1
            if (i==1):
                queation = firebase1.get(keyID,'Queation1')
            elif (i==2):
                queation = firebase1.get(keyID,'Queation2')
            elif (i==3):
                queation = firebase1.get(keyID,'Queation3')
            elif (i==4):
                queation = firebase1.get(keyID,'Queation4')
            elif (i==5):
                queation = firebase1.get(keyID,'Queation5')
            if (queation == 'Excelent'):
                Excelent+=1
            elif (queation == 'Good'):
                Good+=1
            elif (queation == 'Average'):
                Average+=1
            elif (queation == 'Poor'):
                Poor+=1

        Poor = self.percentage(Poor, tf)
        Average = self.percentage(Average, tf)
        Excelent = self.percentage(Excelent, tf)
        Good = self.percentage(Good, tf)
        self.plotPieChart1(Poor, Average, Excelent,Good, tf,i)

        
          

    def removepositive(self,positive_list):
        i=-1
        for f in positive_list:
            i+=1
            j=-1
            for f1 in positive_list:
                j+=1
                if(i!=j):
                    d1 = nlp(f)
                    d2 = nlp(f1)
                    if(d1.similarity(d2)>=0.80):
                        pf=0
                        pf1=0
                        a = TextBlob(f)
                        pf+=a.sentiment.polarity
                        a = TextBlob(f1)
                        pf1+=a.sentiment.polarity
                        if(pf>pf1):
                            positive_list.remove(positive_list[i])
                        elif(pf<=pf1):
                            positive_list.remove(positive_list[j])
                    else:
                        print(d1.similarity(d2))

        print(positive_list)


    def removeneutral(self,neutral_list):
        i=-1
        for f in neutral_list:
            i+=1
            j=-1
            for f1 in neutral_list:
                j+=1
                if(i!=j):
                    d1 = nlp(f)
                    d2 = nlp(f1)
                    if(d1.similarity(d2)>=0.80):
                        neutral_list.remove(neutral_list[i])
                    else:
                        print(d1.similarity(d2))

        print(neutral_list)

        
    def removenegative(self,negative_list):
        i=-1
        for f in negative_list:
            i+=1
            j=-1
            for f1 in negative_list:
                j+=1
                if(i!=j):
                    d1 = nlp(f)
                    d2 = nlp(f1)
                    if(d1.similarity(d2)>=0.80):
                        pf=0
                        pf1=0
                        a = TextBlob(f)
                        pf+=a.sentiment.polarity
                        a = TextBlob(f1)
                        pf1+=a.sentiment.polarity
                        if(pf<pf1):
                            positive_list.remove(negative_list[i])
                        elif(pf>=pf1):
                            positive_list.remove(negative_list[j])
                    else:
                        print(d1.similarity(d2))

        print(negative_list)
                    
    def percentage(self, part, whole):
        temp = 100 * float(part) / float(whole)
        return format(temp, '.2f')

    def plotPieChart(self, positive, negative, neutral, totalfeedback):
        labels = ['Positive [' + str(positive) + '%]', 'Neutral [' + str(neutral) + '%]', 'Negative [' + str(negative) + '%]']
        sizes = [positive, neutral, negative]
        colors = ['lightgreen','gold','lightsalmon']
        patches, texts = plt.pie(sizes, colors=colors, startangle=90)
        plt.legend(patches, labels, loc="best")
        plt.title('How people are reacting on by analyzing ' + str(totalfeedback) + ' feedback.')
        plt.axis('equal')
        plt.tight_layout()
        plt.savefig("finalchart.PNG",bbox_inches="tight",pad_inches=1)



    def plotPieChart1(self,Poor, Average, Excelent,Good, totalfeedback,i):
        labels = ['Poor [' + str(Poor) + '%]', 'Average [' + str(Average) + '%]', 'Excelent [' + str(Excelent) + '%]', 'Good [' + str(Good) + '%]']
        sizes = [Poor, Average, Excelent,Good]
        colors = ['lightgreen','gold','lightsalmon','red']
        patches, texts = plt.pie(sizes, colors=colors, startangle=90)
        plt.legend(patches, labels, loc="best")
        plt.axis('equal')
        plt.tight_layout()
        if (i==1):
                plt.title('1. The content of Expert Talk/Lecture by analyzing' + str(totalfeedback) + ' feedback.')
                plt.savefig("Queation1.PNG",bbox_inches="tight",pad_inches=1)
        elif (i==2):
                plt.title('2. Method of Presentation  by analyzing ' + str(totalfeedback) + ' feedback.')
                plt.savefig("Queation2.PNG",bbox_inches="tight",pad_inches=1)
        elif (i==3):
                plt.title('3. Response to Questions/Queries by Expert by analyzing ' + str(totalfeedback) + ' feedback.')
                plt.savefig("Queation3.PNG",bbox_inches="tight",pad_inches=1)
        elif (i==4):
                plt.title('4. Innovativeness in the Expert Talk/Lecture by analyzing ' + str(totalfeedback) + ' feedback.')
                plt.savefig("Queation4.PNG",bbox_inches="tight",pad_inches=1)
        elif (i==5):
                plt.title('5. Overall Impact of the Expert Talk/Lecture by analyzing ' + str(totalfeedback) + ' feedback.')
                plt.savefig("Queation5.PNG",bbox_inches="tight",pad_inches=1)
            







    
Sentiment_Analysis()
Sentiment_Analysis().fatchdata()

document = Document()
document.add_heading('Feedback Analysis Report..', 0)
document.add_paragraph('Expert_Name :'+Expert_Name, style='Subtitle')
document.add_paragraph('Subject_Name :'+Subject_Name, style='Subtitle')
document.add_paragraph('Date :'+date, style='Subtitle')
document.add_heading('Feed-backs :', level=1)
document.add_paragraph('Positive Feedback:', style='Subtitle')
for i in positive_list:
    document.add_paragraph(
        i, style='List Bullet'
    )
document.add_paragraph('Neutral Feedback:', style='Subtitle')
for i in neutral_list:
    document.add_paragraph(
        i, style='List Bullet'
    )
document.add_paragraph('Negative Feedback:', style='Subtitle')
for i in negative_list:
    document.add_paragraph(
        i, style='List Bullet'
    )
document.add_heading('Final Chart :', level=1)
document.add_picture('finalchart.png', width=Inches(5.00))

document.add_heading('Queation-1  :', level=1)
document.add_picture('Queation1.png', width=Inches(5.00))

document.add_heading('Queation-2  :', level=1)
document.add_picture('Queation2.png', width=Inches(5.00))

document.add_heading('Queation-3  :', level=1)
document.add_picture('Queation3.png', width=Inches(5.00))

document.add_heading('Queation-4  :', level=1)
document.add_picture('Queation4.png', width=Inches(5.00))

document.add_heading('Queation-5  :', level=1)
document.add_picture('Queation5.png', width=Inches(5.00))

document.save('demo3.docx')





